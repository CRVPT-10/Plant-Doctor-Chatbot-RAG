import os
from typing import List
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
from utils.logger import get_logger
from utils.helpers import detect_language

logger = get_logger("document_loader")

class DocumentLoader:
    """Loads documents of various formats (PDF, DOCX, TXT, MD) and returns LangChain Documents."""
    
    @staticmethod
    def load_file(file_path: str) -> List[Document]:
        """
        Loads a single document file and returns a list of LangChain Document objects.
        For PDFs, splits into one Document per page.
        For others, returns a single Document with page=1.
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        documents: List[Document] = []
        filename = os.path.basename(file_path)
        
        try:
            if ext == ".pdf":
                documents = DocumentLoader._load_pdf(file_path, filename)
            elif ext == ".docx":
                documents = DocumentLoader._load_docx(file_path, filename)
            elif ext in [".txt", ".md"]:
                documents = DocumentLoader._load_txt_or_md(file_path, filename)
            else:
                logger.warning(f"Unsupported file type: {ext} for file {file_path}")
                raise ValueError(f"Unsupported file type: {ext}")
                
            logger.info(f"Successfully loaded {len(documents)} pages/documents from {file_path}")
            return documents
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {str(e)}")
            raise e

    @staticmethod
    def _load_pdf(file_path: str, filename: str) -> List[Document]:
        documents: List[Document] = []
        reader = PdfReader(file_path)
        
        # Read page by page to keep track of page numbers
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                lang = detect_language(text)
                metadata = {
                    "source": filename,
                    "page": page_idx + 1,
                    "file_path": file_path,
                    "language": lang,
                    "doc_id": filename
                }
                documents.append(Document(page_content=text, metadata=metadata))
        return documents

    @staticmethod
    def _load_docx(file_path: str, filename: str) -> List[Document]:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text = "\n".join(full_text).strip()
        if not text:
            return []
            
        lang = detect_language(text)
        metadata = {
            "source": filename,
            "page": 1,
            "file_path": file_path,
            "language": lang,
            "doc_id": filename
        }
        return [Document(page_content=text, metadata=metadata)]

    @staticmethod
    def _load_txt_or_md(file_path: str, filename: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
            
        if not text:
            return []
            
        lang = detect_language(text)
        metadata = {
            "source": filename,
            "page": 1,
            "file_path": file_path,
            "language": lang,
            "doc_id": filename
        }
        return [Document(page_content=text, metadata=metadata)]

    @staticmethod
    def load_directory(directory_path: str, recursive: bool = True) -> List[Document]:
        """
        Recursively loads all PDF, DOCX, TXT, MD files in directory.
        """
        documents: List[Document] = []
        if not os.path.exists(directory_path):
            logger.warning(f"Directory does not exist: {directory_path}")
            return []
            
        for root, _, files in os.walk(directory_path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in [".pdf", ".docx", ".txt", ".md"]:
                    file_path = os.path.join(root, file)
                    try:
                        documents.extend(DocumentLoader.load_file(file_path))
                    except Exception as e:
                        logger.error(f"Skipping file {file_path} due to error: {e}")
            if not recursive:
                break
                
        return documents
